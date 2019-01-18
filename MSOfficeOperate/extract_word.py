#!/usr/bin/python3

import os
import sys

from collections import defaultdict

from docx.api import Document

from write_to_file import *

#dates_dict = defaultdict(list)
g_data = dict()

def generate_case(row):
    case = row[2].text + '；' + row[3].text
    case = case.replace('\n','').replace('\r','')
    if not case.endswith('。'):
        case += '。'
    return case

def load_data():
    print("load data ...")
    document = Document('ats-2-9.docx')

    f_name = ''
    table_index = 0
    for paragraph in document.paragraphs:
        if paragraph.style.name == "Heading 2":
            # feature
            f_name = paragraph.text
            g_data[f_name] = defaultdict(list)
        elif paragraph.style.name == "Heading 3":
            # test
            g_data[f_name][paragraph.text] = list()
            # case
            table = document.tables[table_index]
            for i, row in enumerate(table.rows):
                if i != 0:
                    g_data[f_name][paragraph.text].append(generate_case(row.cells))
            table_index += 1



#    table = doc.tables[0]
#    for i, row in enumerate(table.rows):
#        for cell in row.cells:
#            print(cell.text)



def main():
    load_data()
    print(g_data)
    write_data('ats-2-9', g_data)


if __name__ == "__main__":
    main()
